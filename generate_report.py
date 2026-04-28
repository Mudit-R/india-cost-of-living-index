"""
Generate a presentable DOCX technical report using python-docx.
Run from the project root with the venv active.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette (hex strings for bg, RGBColor for font) ───────────────
DARK_BLUE_HEX  = "1F497D"
MID_BLUE_HEX   = "2E75B6"
LIGHT_BLUE_HEX = "D6E4F0"
ALT_ROW_HEX    = "F2F7FB"
WHITE_HEX      = "FFFFFF"

DARK_BLUE  = RGBColor(0x1F, 0x49, 0x7D)
MID_BLUE   = RGBColor(0x2E, 0x75, 0xB6)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GREY  = RGBColor(0x40, 0x40, 0x40)

# ── Helpers ───────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="2E75B6", sz="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def para_space(doc, pts_before=0, pts_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(pts_before)
    p.paragraph_format.space_after  = Pt(pts_after)
    return p

def add_horizontal_rule(doc, color="2E75B6"):
    p  = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

def add_heading(doc, text, level=1):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        run.font.size  = Pt(16)
        run.font.bold  = True
        run.font.color.rgb = DARK_BLUE
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(4)
        add_horizontal_rule(doc)
    elif level == 2:
        run.font.size  = Pt(13)
        run.font.bold  = True
        run.font.color.rgb = MID_BLUE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(3)
    elif level == 3:
        run.font.size  = Pt(11.5)
        run.font.bold  = True
        run.font.color.rgb = DARK_GREY
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
    return p

def add_body(doc, text, italic=False):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.color.rgb = DARK_GREY
    run.font.italic    = italic
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(15)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GREY
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)

def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'EFEFEF')
    p._p.get_or_add_pPr().append(shd)
    p.paragraph_format.left_indent   = Cm(0.5)
    p.paragraph_format.space_before  = Pt(4)
    p.paragraph_format.space_after   = Pt(4)
    return p

def style_table(table, headers, rows, col_widths=None):
    """Add a styled table with blue header row and alternating row colours."""
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style     = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run  = cell.paragraphs[0].runs[0]
        run.font.bold       = True
        run.font.size       = Pt(10.5)
        run.font.color.rgb  = WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, DARK_BLUE_HEX)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        bg  = ALT_ROW_HEX if ri % 2 == 0 else WHITE_HEX
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            run  = cell.paragraphs[0].runs[0]
            run.font.size      = Pt(10)
            run.font.color.rgb = DARK_GREY
            set_cell_bg(cell, bg)

# ── Create document ───────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# Default font
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

# ── Header ────────────────────────────────────────────────────────────────
header = doc.sections[0].header
header.paragraphs[0].clear()
hrun = header.paragraphs[0].add_run("Cost of Living Index — Indian Cities   |   Technical Report")
hrun.font.size      = Pt(9)
hrun.font.color.rgb = MID_BLUE
hrun.font.italic    = True
header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_horizontal_rule(doc.sections[0].header)

# ── Footer ────────────────────────────────────────────────────────────────
footer = doc.sections[0].footer
fpar   = footer.paragraphs[0]
fpar.clear()
frun = fpar.add_run("BMP Data Modelling  ·  April 2026  ·  Confidential")
frun.font.size      = Pt(8.5)
frun.font.color.rgb = RGBColor(0x90, 0x90, 0x90)
fpar.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ─────────────────────────────────────────────────────────────────────────
# TITLE PAGE
# ─────────────────────────────────────────────────────────────────────────
para_space(doc, pts_before=60)

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("Cost of Living Index")
tr.font.size      = Pt(30)
tr.font.bold      = True
tr.font.color.rgb = DARK_BLUE

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run("A Data-Driven Analysis of 50 Indian Cities")
sr.font.size      = Pt(16)
sr.font.color.rgb = MID_BLUE
sr.font.italic    = True

para_space(doc, pts_before=8)
add_horizontal_rule(doc)
para_space(doc, pts_before=8)

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta_p.add_run("Technical Report  ·  BMP Data Modelling  ·  April 2026")
mr.font.size      = Pt(12)
mr.font.color.rgb = DARK_GREY

para_space(doc, pts_before=30)

# Summary box as a single-cell table
sumbox = doc.add_table(rows=1, cols=1)
sumbox.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = sumbox.cell(0, 0)
set_cell_bg(cell, LIGHT_BLUE_HEX)
cell.width = Cm(14)
kv = [
    ("Cities covered",    "50"),
    ("Data files processed", "94+"),
    ("Expense categories", "8"),
    ("Base city",         "Delhi (Index = 100)"),
    ("Language",          "Python 3.8+"),
]
for k, v in kv:
    p = cell.add_paragraph()
    r1 = p.add_run(f"{k}:  ")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = DARK_BLUE
    r2 = p.add_run(v)
    r2.font.size = Pt(11)
    r2.font.color.rgb = DARK_GREY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.4)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────
# SECTION 1 — INTRODUCTION
# ─────────────────────────────────────────────────────────────────────────
add_heading(doc, "1.  Introduction")
add_body(doc,
    "Comparing the cost of living across cities is rarely straightforward. Prices vary not just "
    "between metros and smaller towns but also within the same tier, driven by differences in "
    "local supply chains, regulatory environments, and consumer habits. Most existing indices "
    "either cover too few cities, rely on outdated survey data, or reduce everything to a single "
    "number without showing what is driving costs in any particular place.")
add_body(doc,
    "This report documents the design, data pipeline, and results of a cost of living index "
    "built for 50 Indian cities. The index is constructed entirely from raw transactional data "
    "pulled from live platforms — property listings, grocery delivery apps, ride-hailing "
    "services, tutor marketplaces, and restaurant aggregators. Every step of the pipeline, from "
    "data cleaning to missing value handling, is described here so the work can be reproduced "
    "or extended.")
add_body(doc,
    "Delhi is used as the base city, assigned an index of 100. All other cities are expressed "
    "relative to that baseline. A city at 130 costs 30% more than Delhi in aggregate; "
    "a city at 80 costs 20% less.")

# ─────────────────────────────────────────────────────────────────────────
# SECTION 2 — SCOPE
# ─────────────────────────────────────────────────────────────────────────
add_heading(doc, "2.  Scope and Objectives")
add_body(doc, "The primary goals of this project were:")
for pt in [
    "Build a reproducible, city-level cost of living index covering 50 cities across different tiers of Indian urban settlement.",
    "Use real transactional data rather than consumer surveys, which are expensive to run and quick to date.",
    "Weight each expense category according to how much of a typical urban household budget it actually consumes.",
    "Generate visualisations that allow comparisons across individual categories, not just at the aggregate level.",
    "Identify where the largest cost gaps between cities occur and which components drive them.",
    "Connect the index to a practical recommendation engine for individuals considering relocation.",
]:
    add_bullet(doc, pt)

# ─────────────────────────────────────────────────────────────────────────
# SECTION 3 — DATA SOURCES
# ─────────────────────────────────────────────────────────────────────────
add_heading(doc, "3.  Data Sources")
add_body(doc,
    "Data was collected from seven categories, each representing a distinct slice of household "
    "spending. The table below provides a quick overview before each source is described in detail.")

# Overview table
t = doc.add_table(rows=1, cols=4)
style_table(t,
    headers=["Category", "Source", "Format", "Coverage"],
    rows=[
        ["Housing",       "MagicBricks",      "CSV (50 files)", "50 / 50 cities"],
        ["Grocery",       "Blinkit",          "Excel (41 files)", "41 / 50 cities"],
        ["Transport",     "Uber / Fuel data", "Excel + CSV",    "42 / 50 (Uber)"],
        ["Healthcare",    "Physician fee data", "Excel (wide)",  "50 / 50 cities"],
        ["Education",     "Tutor marketplace","Excel",          "50 / 50 cities"],
        ["Electricity",   "State tariff data","Excel",          "50 / 50 cities"],
        ["Entertainment", "Swiggy + cinemas", "Excel (2 files)","50 / 50 cities"],
    ],
    col_widths=[3, 3.5, 3, 3]
)
para_space(doc, pts_before=6)

add_heading(doc, "3.1  Housing — MagicBricks", level=2)
add_body(doc,
    "Property listing data was collected from MagicBricks for 50 cities. Each city has its own "
    "CSV file named CityName_magicbricks.csv. The files contain individual property listings with "
    "city name, number of floors, number of bathrooms, carpet area in square feet, listed price, "
    "and derived price per square foot. The metric used is price per square foot rather than the "
    "absolute price, which makes comparisons more consistent given that the size distribution of "
    "properties varies considerably between cities.")

t2 = doc.add_table(rows=1, cols=2)
style_table(t2,
    headers=["City", "Median Price / sq ft (INR)"],
    rows=[
        ["Delhi",      "10,231"],
        ["Mumbai",     "23,619"],
        ["Bengaluru",  "11,029"],
        ["Jaipur",     "3,579"],
        ["Malappuram", "1,512"],
    ]
)
para_space(doc, pts_before=6)

add_heading(doc, "3.2  Grocery — Blinkit", level=2)
add_body(doc,
    "Grocery price data was sourced from Blinkit, an online grocery delivery platform with "
    "strong city-level penetration across tier-1 and tier-2 markets. Files are in Excel format, "
    "one per city, and contain product names alongside current listed prices. The IQR-cleaned "
    "median price across all products in each city is used as the city-level metric. Coverage "
    "is 41 out of 50 cities; the remaining nine are filled with the national median (INR 42).")

add_heading(doc, "3.3  Transport — Uber and Fuel Prices", level=2)
add_body(doc,
    "Transport costs are captured through two files: an Excel sheet with Uber's per-kilometre "
    "rate by city, and a CSV with petrol, diesel, CNG, and LPG prices. Uber pricing is used "
    "as the index metric because it represents on-demand urban commuting. Fuel prices are "
    "available in the dataset but are not directly weighted in the index, since price variation "
    "across cities is driven mainly by state-level taxes rather than local market conditions.")

add_heading(doc, "3.4  Healthcare — General Physician Fees", level=2)
add_body(doc,
    "Doctor consultation fees were compiled in a wide-format Excel file where each row "
    "corresponds to a city and each subsequent column holds a fee observation from a different "
    "clinic or practitioner. Only general physician fees are included. Specialist fees — which "
    "can be an order of magnitude higher — are excluded because they do not represent routine "
    "household spending.")

add_heading(doc, "3.5  Education — Tutor Marketplace", level=2)
add_body(doc,
    "Education cost data comes from a tutor listing platform containing 60,256 listings with "
    "tutor ID, city, and a fee range string such as 'INR 300–1,000/hour'. Monthly fee listings "
    "are excluded — only per-hour rates are retained (46,965 records). "
    "Fee extraction uses a regular expression; for a range, the midpoint is used. "
    "The per-city metric is the mean of IQR-cleaned hourly rates. Cities average around "
    "940 listings each, giving a stable estimate of local tutoring costs.")

add_heading(doc, "3.6  Utilities — Electricity Tariffs", level=2)
add_body(doc,
    "Electricity rates come from a single Excel file with one row per city, containing the "
    "effective tariff in rupees per kilowatt-hour sourced from state electricity board schedules. "
    "The range across 50 cities is INR 3.00 to 8.50 per unit — variation is explained entirely "
    "by differences in state tariff structures.")

add_heading(doc, "3.7  Entertainment — Movies and Restaurants", level=2)
add_body(doc,
    "Movie ticket prices were gathered for 5–15 cinemas per city. Restaurant prices come from "
    "a Swiggy dataset listing average meal costs per location. Swiggy prices are listed 'for two', "
    "so all values are halved before analysis. City name aliases require handling for this "
    "dataset — Bangalore maps to Bengaluru, Mysore to Mysuru, Trichy to Tiruchirappalli, and "
    "so on. Without this normalisation, city-level joins between datasets fail silently.")

# ─────────────────────────────────────────────────────────────────────────
# SECTION 4 — METHODOLOGY
# ─────────────────────────────────────────────────────────────────────────
add_heading(doc, "4.  Methodology")

add_heading(doc, "4.1  Outlier Removal — IQR Method", level=2)
add_body(doc,
    "Before any aggregation, each multi-observation dataset goes through IQR-based outlier "
    "removal at the city level. The valid price range is defined as:")
add_code(doc, "Valid range  =  [ Q1 − 1.5 × IQR ,  Q3 + 1.5 × IQR ]")
add_body(doc,
    "This removes extreme observations — luxury penthouses from housing data, IMAX ticket "
    "prices from movie data, premium specialist tutors from education data — without requiring "
    "manual curation. The assumption is that the index should reflect what a typical resident "
    "actually pays, not the most expensive option available.")

add_heading(doc, "4.2  City Name Normalisation", level=2)
add_body(doc,
    "A mapping dictionary handles known aliases and spelling variations before any dataset "
    "joins are performed:")

t3 = doc.add_table(rows=1, cols=2)
style_table(t3,
    headers=["Alias in Raw Data", "Canonical Name"],
    rows=[
        ["Bangalore",  "Bengaluru"],
        ["Bombay",     "Mumbai"],
        ["Calcutta",   "Kolkata"],
        ["Mysore",     "Mysuru"],
        ["Trichy",     "Tiruchirappalli"],
        ["Trivandrum", "Thiruvananthapuram"],
        ["Merrut",     "Meerut"],
    ]
)
para_space(doc, pts_before=6)

add_heading(doc, "4.3  Index Calculation", level=2)
add_body(doc, "Each component index is expressed relative to Delhi:")
add_code(doc, "Component Index (city)  =  ( City Price  /  Delhi Price )  ×  100")
add_body(doc, "The overall cost of living index is a weighted average of the eight component indices:")
add_code(doc, "Cost Index  =  Σ  weight_i  ×  Component Index_i")

add_heading(doc, "4.4  Weight Derivation", level=2)
add_body(doc,
    "Raw weights are assigned based on typical urban household expenditure patterns in India "
    "and re-normalised to sum to 100% at runtime.")

t4 = doc.add_table(rows=1, cols=4)
style_table(t4,
    headers=["Component", "Raw Weight", "Normalised", "Basis"],
    rows=[
        ["Grocery",     "0.30",   "36.36%", "Daily consumption — largest recurring expense"],
        ["Housing",     "0.25",   "30.30%", "Rent or EMI — largest fixed monthly cost"],
        ["Transport",   "0.09",   "10.91%", "Daily commute via ride-hailing"],
        ["Healthcare",  "0.053",  "6.42%",  "Routine doctor visits"],
        ["Education",   "0.050",  "6.06%",  "Private tutoring, widespread in urban India"],
        ["Restaurant",  "0.040",  "4.85%",  "Meals out and food delivery"],
        ["Electricity", "0.025",  "3.03%",  "Monthly utility billing"],
        ["Movies",      "0.0171", "2.07%",  "Discretionary entertainment proxy"],
        ["Total",       "0.8251", "100%",   ""],
    ]
)
para_space(doc, pts_before=6)

add_heading(doc, "4.5  Missing Value Handling", level=2)
add_body(doc,
    "When a city has no data for a given component, the gap is filled with the median value "
    "across all cities with available data. Median is preferred over mean because the "
    "distribution of component values across cities tends to be right-skewed — a few "
    "high-cost cities would pull the mean upward, overstating the typical national value.")

t5 = doc.add_table(rows=1, cols=3)
style_table(t5,
    headers=["Component", "Cities Missing", "Fill Value"],
    rows=[
        ["Uber (transport)", "8",  "INR 18.91 / km"],
        ["Grocery",          "9",  "INR 42.00"],
        ["Housing",          "0",  "No imputation needed"],
    ]
)
para_space(doc, pts_before=6)

# ─────────────────────────────────────────────────────────────────────────
# SECTION 5 — RESULTS
# ─────────────────────────────────────────────────────────────────────────
add_heading(doc, "5.  Results")

add_heading(doc, "5.1  Full City Rankings", level=2)
add_body(doc,
    "The table below gives the complete cost of living index for all 50 cities. Cities above "
    "100 are more expensive than Delhi; cities below 100 are more affordable.")

ranking_rows = [
    ["1",  "Mumbai",             "162.79", "Expensive"],
    ["2",  "Bengaluru",          "114.76", "Expensive"],
    ["3",  "Kozhikode",          "107.66", "Expensive"],
    ["4",  "Hyderabad",          "107.09", "Expensive"],
    ["5",  "Patna",              "100.72", "Near Baseline"],
    ["6",  "Delhi",              "100.00", "Base City"],
    ["7",  "Pune",               "99.61",  "Moderate"],
    ["8",  "Kolkata",            "98.57",  "Moderate"],
    ["9",  "Ahmedabad",          "95.42",  "Moderate"],
    ["10", "Coimbatore",         "93.04",  "Moderate"],
    ["11", "Chennai",            "92.25",  "Moderate"],
    ["12", "Raipur",             "89.95",  "Moderate"],
    ["13", "Bhubaneswar",        "89.18",  "Moderate"],
    ["14", "Jaipur",             "88.38",  "Moderate"],
    ["15", "Rajkot",             "85.78",  "Moderate"],
    ["16", "Bhopal",             "85.25",  "Moderate"],
    ["17", "Jabalpur",           "84.32",  "Moderate"],
    ["18", "Sangli",             "83.17",  "Moderate"],
    ["19", "Agra",               "82.75",  "Moderate"],
    ["20", "Thrissur",           "82.52",  "Moderate"],
    ["21", "Salem",              "81.95",  "Affordable"],
    ["22", "Mysuru",             "81.08",  "Affordable"],
    ["23", "Kollam",             "81.05",  "Affordable"],
    ["24", "Nagpur",             "81.03",  "Affordable"],
    ["25", "Kolhapur",           "80.48",  "Affordable"],
    ["26", "Lucknow",            "80.41",  "Affordable"],
    ["27", "Erode",              "80.35",  "Affordable"],
    ["28", "Chandigarh",         "80.16",  "Affordable"],
    ["29", "Aurangabad",         "79.87",  "Affordable"],
    ["30", "Visakhapatnam",      "79.52",  "Affordable"],
    ["31", "Kochi",              "79.24",  "Affordable"],
    ["32", "Kanpur",             "79.21",  "Affordable"],
    ["33", "Asansol",            "78.39",  "Affordable"],
    ["34", "Vadodara",           "77.82",  "Affordable"],
    ["35", "Madurai",            "77.63",  "Affordable"],
    ["36", "Meerut",             "77.60",  "Affordable"],
    ["37", "Kottayam",           "77.11",  "Affordable"],
    ["38", "Mangaluru",          "76.59",  "Affordable"],
    ["39", "Surat",              "76.49",  "Affordable"],
    ["40", "Thiruvananthapuram", "76.36",  "Affordable"],
    ["41", "Indore",             "76.10",  "Affordable"],
    ["42", "Nashik",             "75.71",  "Affordable"],
    ["43", "Amaravati",          "75.65",  "Affordable"],
    ["44", "Hubli",              "75.10",  "Affordable"],
    ["45", "Kannur",             "74.34",  "Affordable"],
    ["46", "Jamnagar",           "72.77",  "Affordable"],
    ["47", "Ludhiana",           "72.50",  "Affordable"],
    ["48", "Tiruchirappalli",    "71.41",  "Affordable"],
    ["49", "Malappuram",         "69.54",  "Very Affordable"],
    ["50", "Solapur",            "68.20",  "Very Affordable"],
]
t6 = doc.add_table(rows=1, cols=4)
style_table(t6, headers=["Rank", "City", "Index", "Category"], rows=ranking_rows)
para_space(doc, pts_before=6)

add_heading(doc, "5.2  Component-Level Variation", level=2)
add_body(doc,
    "The aggregate index can mask substantial differences at the component level. Housing is "
    "the most volatile — Mumbai's housing index of 230.86 is 15 times that of Malappuram "
    "(14.77), a spread far larger than any other category. Transport also shows wide variation "
    "because Uber's per-km rates differ significantly by city: Hyderabad at 242.31 versus "
    "Mangaluru at 43.75.")

t7 = doc.add_table(rows=1, cols=3)
style_table(t7,
    headers=["Component", "Highest City (Index)", "Lowest City (Index)"],
    rows=[
        ["Housing",     "Mumbai (230.86)",    "Malappuram (14.77)"],
        ["Grocery",     "Kozhikode (151.21)", "Nashik (90.94)"],
        ["Transport",   "Hyderabad (242.31)", "Mangaluru (43.75)"],
        ["Healthcare",  "Mumbai (135.0)",     "Salem (15.0)"],
        ["Education",   "Mumbai (113.87)",    "Salem (15.38)"],
        ["Electricity", "Mumbai (212.5)",     "Visakhapatnam (75.0)"],
        ["Restaurant",  "Mumbai (120.0)",     "Kolkata (80.0)"],
        ["Movies",      "Kozhikode (153.85)", "Salem / Madurai (15.38)"],
    ]
)
para_space(doc, pts_before=6)

add_heading(doc, "5.3  Regional Patterns", level=2)
add_body(doc,
    "Kerala cities present an unusual picture. Despite being relatively affordable in housing, "
    "Kozhikode registers a grocery index of 151.21 — the highest in the dataset. This is "
    "consistent with Kerala's dependence on food imports from neighbouring states and its "
    "generally higher grocery price level.")
add_body(doc,
    "Gujarat cities — Rajkot, Surat, Ahmedabad, Vadodara, Jamnagar — cluster in the "
    "moderate-to-affordable range across most components, consistent with the state's "
    "reputation for cost-effective urban living despite strong economic activity.")
add_body(doc,
    "Maharashtra presents a split picture: Mumbai is in a class of its own while Pune, Nagpur, "
    "Nashik, Aurangabad, Kolhapur, and Sangli all sit comfortably in the affordable-to-moderate "
    "band. Proximity to Mumbai does not appear to push costs significantly higher in these cities.")

# ─────────────────────────────────────────────────────────────────────────
# SECTION 6 — IMPLEMENTATION
# ─────────────────────────────────────────────────────────────────────────
add_heading(doc, "6.  Implementation")

add_heading(doc, "6.1  Code Structure", level=2)
t8 = doc.add_table(rows=1, cols=2)
style_table(t8,
    headers=["Module", "Responsibility"],
    rows=[
        ["data_loader.py",       "Load and clean each data source; return standardised DataFrames"],
        ["cost_calculator.py",   "Apply weights, normalise, and compute the index"],
        ["visualizer.py",        "Generate all 23 output charts"],
        ["ml_classification.py", "Classify cities into cost tiers using ML models"],
        ["weight_optimizer.py",  "Search for weight combinations that minimise a cost function"],
        ["recommender.py",       "Priority-based city recommendation logic"],
        ["main.py",              "Top-level orchestration; runs the full pipeline end to end"],
    ]
)
para_space(doc, pts_before=6)

add_heading(doc, "6.2  Running the Pipeline", level=2)
add_code(doc, "# Full pipeline from the project root\n./run.sh\n\n# Equivalent manual run\ncd src && python3 main.py\n\n# Interactive city recommender\nstreamlit run website/app.py")

add_heading(doc, "6.3  Dependencies", level=2)
t9 = doc.add_table(rows=1, cols=2)
style_table(t9,
    headers=["Library", "Purpose"],
    rows=[
        ["pandas",       "Data loading, merging, and transformation"],
        ["numpy",        "Array operations and IQR calculations"],
        ["matplotlib",   "All chart rendering"],
        ["seaborn",      "Heatmap and styled statistical charts"],
        ["openpyxl",     "Reading Excel files (.xlsx)"],
        ["streamlit",    "Web interface for city recommender"],
        ["scikit-learn", "ML classification models"],
    ]
)
para_space(doc, pts_before=6)

# ─────────────────────────────────────────────────────────────────────────
# SECTION 7 — ML
# ─────────────────────────────────────────────────────────────────────────
add_heading(doc, "7.  Machine Learning Component")

add_heading(doc, "7.1  City Tier Classification", level=2)
add_body(doc,
    "A classification model is trained to assign cities to one of three tiers — Expensive, "
    "Moderate, and Affordable — using the eight component indices as features. Four algorithms "
    "are compared: Random Forest, Support Vector Machine, Decision Tree, and Logistic Regression. "
    "Random Forest and SVM generally outperform the simpler classifiers on this dataset given "
    "the non-linear relationships between components. Training on 50 data points is a known "
    "limitation; the models are illustrative rather than production-grade.")

add_heading(doc, "7.2  Recommendation Engine", level=2)
add_body(doc,
    "The city recommender takes user-defined priorities for each cost category and scores all "
    "50 cities against those preferences. Cities where the actual component index aligns with "
    "the stated preference score higher. This is surfaced through a Streamlit web interface "
    "that shows the full component breakdown for each recommended city.")

add_heading(doc, "7.3  Weight Optimisation", level=2)
add_body(doc,
    "weight_optimizer.py uses Scipy's optimisation routines to find the component weight set "
    "that minimises variance in city rankings when individual components are perturbed — a "
    "measure of index robustness. The optimised weights are compared to the original "
    "expenditure-based weights to assess how sensitive the rankings are to the weighting scheme.")

# ─────────────────────────────────────────────────────────────────────────
# SECTION 8 — LIMITATIONS
# ─────────────────────────────────────────────────────────────────────────
add_heading(doc, "8.  Limitations")

for lim in [
    ("Data is not contemporaneous.",
     "The various datasets were collected at different points in time. Grocery prices from "
     "Blinkit may be from a different quarter than the MagicBricks listings. The index is a "
     "snapshot, not a continuous tracker."),
    ("Quality differences are not captured.",
     "An average apartment in Bengaluru at INR 11,000 per sq ft is almost certainly better "
     "located than one at the same price in a smaller city. The index makes no adjustment for "
     "quality."),
    ("Grocery imputation affects nine cities.",
     "The nine cities without Blinkit coverage are assigned the national median grocery value, "
     "fixing their grocery index at 100. Comparisons involving those cities on the grocery "
     "dimension should be treated with caution."),
    ("Small sample for ML.",
     "Fifty cities is not a large training corpus. The classification models are illustrative "
     "rather than production-grade."),
    ("Income data is absent.",
     "Without income data, the index cannot produce an affordability measure — a ratio of cost "
     "to local salary that would be more useful for relocation decisions."),
]:
    p = doc.add_paragraph()
    r1 = p.add_run(lim[0] + "  ")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = DARK_BLUE
    r2 = p.add_run(lim[1])
    r2.font.size = Pt(11)
    r2.font.color.rgb = DARK_GREY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(15)

# ─────────────────────────────────────────────────────────────────────────
# SECTION 9 — CONCLUSIONS
# ─────────────────────────────────────────────────────────────────────────
add_heading(doc, "9.  Conclusions")
add_body(doc,
    "The cost of living gap between Indian cities is real and large. Mumbai's index of 162.79 "
    "is more than twice Solapur's 68.20, and that ratio holds even after removing extreme values. "
    "Housing is the primary driver — it accounts for 30% of the index weight and shows the "
    "widest variation of any component.")
add_body(doc,
    "Tier-2 cities are not uniformly affordable. Some, like Coimbatore, Bhubaneswar, and Jaipur, "
    "sit only modestly below Delhi. Others, like Malappuram, Jamnagar, and Tiruchirappalli, are "
    "genuinely low-cost environments. The index makes it easier to distinguish between these "
    "two groups than headline city-tier labels do.")
add_body(doc,
    "The pipeline is modular and extensible: adding a new data source requires writing a new "
    "loader function in data_loader.py, adding the corresponding weight and label in the "
    "configuration, and re-running. The next logical step would be to put this on a time series "
    "footing — running the pipeline quarterly — so that cost trends can be tracked over time "
    "rather than read as a single snapshot.")

# ─────────────────────────────────────────────────────────────────────────
# APPENDIX
# ─────────────────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "Appendix A:  Output Files")
t10 = doc.add_table(rows=1, cols=2)
style_table(t10,
    headers=["File", "Description"],
    rows=[
        ["outputs/reports/cost_index_results.csv",               "Full results: raw prices, component indices, overall index"],
        ["outputs/visualizations/cost_of_living_all_cities.png", "Horizontal bar chart, all 50 cities ranked"],
        ["outputs/visualizations/heatmap_all_50_cities.png",     "Component heatmap across all cities and components"],
        ["outputs/visualizations/radar_chart_comparison.png",    "Multi-city radar comparison"],
        ["outputs/visualizations/component_distribution_boxplot.png", "Box plot distribution per component"],
        ["outputs/visualizations/housing_vs_overall.png",        "Housing index vs. overall index scatter"],
        ["outputs/visualizations/ml_algorithm_comparison.png",   "ML classifier accuracy comparison"],
        ["outputs/visualizations/ml_confusion_matrix.png",       "ML confusion matrix for best classifier"],
        ["outputs/visualizations/weight_optimization_comparison.png", "Weight sensitivity analysis"],
        ["outputs/visualizations/top_bottom_cities.png",         "Top 5 vs. bottom 5 cities comparison"],
    ]
)

para_space(doc, pts_before=10)
add_heading(doc, "Appendix B:  Full City List")
add_body(doc,
    "Agra · Ahmedabad · Amaravati · Asansol · Aurangabad · Bengaluru · Bhopal · Bhubaneswar · "
    "Chandigarh · Chennai · Coimbatore · Delhi · Erode · Hubli · Hyderabad · Indore · "
    "Jabalpur · Jaipur · Jamnagar · Kannur · Kanpur · Kochi · Kolhapur · Kolkata · Kollam · "
    "Kottayam · Kozhikode · Lucknow · Ludhiana · Madurai · Malappuram · Mangaluru · Meerut · "
    "Mumbai · Mysuru · Nagpur · Nashik · Patna · Pune · Raipur · Rajkot · Salem · Sangli · "
    "Solapur · Surat · Thiruvananthapuram · Thrissur · Tiruchirappalli · Vadodara · Visakhapatnam")

# ── Save ──────────────────────────────────────────────────────────────────
doc.save("TECHNICAL_REPORT.docx")
print("Saved: TECHNICAL_REPORT.docx")
